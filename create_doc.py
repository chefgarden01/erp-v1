from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ===== 표지 =====
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('그트플 통합관리시스템')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(26, 86, 219)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('공통코드 관리 체계 설계안')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(75, 85, 99)

doc.add_paragraph('')
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('2026년 3월 7일').font.size = Pt(11)

doc.add_page_break()

# ===== 1. 현재 문제점 =====
doc.add_heading('1. 현재 문제점', level=1)

doc.add_paragraph(
    '현재 기초정보 코드 24종의 드롭다운 항목(대분류, 중분류, 상태, 용도 등)이 '
    '프로그램 코드 내부에 고정(하드코딩)되어 있습니다.'
)

problems = [
    ('새 항목 추가 불가', '사용자가 직접 드롭다운 항목을 추가할 수 없음. 개발자가 코드를 수정해야 반영'),
    ('입력 중 추가 불가', '데이터 입력 중 필요한 코드가 없으면 작업을 중단하고 요청해야 함'),
    ('코드 간 불일치', '마스터상품의 "대분류"와 다른 모듈에서 사용하는 분류가 일치하지 않을 수 있음'),
    ('확장성 부족', '상품, 거래처, 자재 등 지속적으로 늘어나는 분류 체계에 대응 불가'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.columns[0].width = Cm(5)
table.columns[1].width = Cm(11)
hdr = table.rows[0].cells
hdr[0].text = '문제'
hdr[1].text = '상세 설명'
for prob, desc in problems:
    row = table.add_row().cells
    row[0].text = prob
    row[1].text = desc

doc.add_paragraph('')

# ===== 2. 해결 방안: 공통코드 관리 체계 =====
doc.add_heading('2. 해결 방안: 공통코드 관리 메뉴', level=1)

doc.add_paragraph(
    '모든 드롭다운 항목을 "공통코드"라는 하나의 관리 체계로 통합합니다. '
    '사용자가 직접 코드를 추가/수정/정지할 수 있으며, 입력 중에도 즉시 새 코드를 생성할 수 있습니다.'
)

doc.add_heading('2-1. 공통코드란?', level=2)
doc.add_paragraph(
    '시스템 전체에서 드롭다운, 분류, 구분값으로 사용되는 모든 항목을 '
    '"그룹코드 → 상세코드" 2단계 구조로 관리하는 체계입니다.'
)

doc.add_heading('2-2. 구조 설계', level=2)

p = doc.add_paragraph()
p.add_run('■ 그룹코드 (상위)').bold = True
doc.add_paragraph('드롭다운이 사용되는 "분류 항목" 자체를 관리', style='List Bullet')
doc.add_paragraph('예: "마스터상품_대분류", "하우스_용도", "거래처_구분" 등', style='List Bullet')

p = doc.add_paragraph()
p.add_run('■ 상세코드 (하위)').bold = True
doc.add_paragraph('그룹코드에 속하는 실제 선택값을 관리', style='List Bullet')
doc.add_paragraph('예: 대분류 → 엽채류, 과채류, 근채류, 과일류 등', style='List Bullet')

# 구조 테이블
doc.add_paragraph('')
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['그룹코드ID', '그룹코드명', '상세코드ID', '상세코드명', '정렬순서']):
    hdr[i].text = h

examples = [
    ('PRD_CAT_L', '마스터상품_대분류', 'PRD_CAT_L_001', '엽채류', '1'),
    ('PRD_CAT_L', '마스터상품_대분류', 'PRD_CAT_L_002', '과채류', '2'),
    ('PRD_CAT_L', '마스터상품_대분류', 'PRD_CAT_L_003', '근채류', '3'),
    ('PRD_CAT_M', '마스터상품_중분류', 'PRD_CAT_M_001', '쌈채소', '1'),
    ('PRD_CAT_M', '마스터상품_중분류', 'PRD_CAT_M_002', '양념채소', '2'),
    ('GHS_USAGE', '하우스_용도', 'GHS_USAGE_001', '생산동', '1'),
    ('GHS_USAGE', '하우스_용도', 'GHS_USAGE_002', '육묘동', '2'),
    ('GHS_STATUS', '하우스_상태', 'GHS_STATUS_001', '생육중', '1'),
    ('CST_CHANNEL', '판매채널', 'CST_CHANNEL_001', 'GTP-오프라인', '1'),
    ('CST_CHANNEL', '판매채널', 'CST_CHANNEL_002', 'CHEF-온라인', '2'),
    ('SUP_TYPE', '매입거래처_구분', 'SUP_TYPE_001', '상품매입', '1'),
    ('COST_CAT', '비용_대분류', 'COST_CAT_001', '인건비', '1'),
]
for ex in examples:
    row = table.add_row().cells
    for i, v in enumerate(ex):
        row[i].text = v

doc.add_paragraph('')

# ===== 3. 현재 하드코딩된 드롭다운 목록 =====
doc.add_heading('3. 공통코드로 전환해야 할 드롭다운 목록', level=1)

doc.add_paragraph(
    '현재 24종 코드에서 사용 중인 모든 드롭다운 항목을 정리한 목록입니다. '
    '이 항목들이 공통코드 관리 메뉴로 이동됩니다.'
)

dropdowns = [
    ('마스터상품', 'PRD_CAT_L', '대분류', '엽채류, 과채류, 근채류, 과일류, 양념채소, 버섯류, 기타'),
    ('마스터상품', 'PRD_CAT_M', '중분류', '쌈채소, 샐러드, 양념채소, 열매채소, 뿌리채소, 기타'),
    ('마스터상품', 'PRD_UNIT', '기본단위', 'kg, g, 개, 박스, 봉, 팩, 묶음'),
    ('마스터상품', 'PRD_ORIGIN', '원산지', '국산, 수입(중국), 수입(미국), 수입(베트남), 수입(기타)'),
    ('마스터상품', 'PRD_CHANNEL', '매입채널', '농장, 시장, 외부'),
    ('마스터상품', 'PRD_STORAGE', '보관구분', '냉장, 냉동, 상온'),
    ('SKU상품', 'SKU_PKG', '포장형태', '봉지, 박스, 팩, 트레이, 기타'),
    ('SKU상품', 'SKU_SHIP', '출고방식', '직배송, 택배, 화물, 퀵'),
    ('포장자재', 'PKG_CAT', '자재분류', '포장팩, 박스, 라벨지, 테이프, 완충재, 밴드, 기타'),
    ('소모품', 'CSM_CAT', '소모품분류', '사무용품, 탕비비품, 청소용품, 위생용품, 기타'),
    ('물류자재', 'EQP_CAT', '물류자재분류', '바구니, 파레트, 진열대, 운반카트, 기타'),
    ('시설/기구', 'FCL_CAT', '시설기구분류', '시설, 기구'),
    ('매출거래처', 'CST_CHANNEL', '판매채널', 'GTP-오프라인, CHEF-온라인'),
    ('매출거래처', 'CST_PAY', '결제조건', '선불, 후불, 월정산, 기타'),
    ('매입거래처', 'SUP_TYPE', '거래처구분', '상품매입, 자재매입, 소모품매입, 물류'),
    ('매입거래처', 'SUP_PAY', '결제조건', '현금, 외상, 월정산, 기타'),
    ('역발행', 'RIV_TYPE', '발행유형', '수수료, 장려금, 판촉비, 기타'),
    ('직원', 'EMP_TYPE', '직원구분', '관리직, 현장직, 계약직'),
    ('직원', 'EMP_NAT', '국적', '한국, 캄보디아, 기타'),
    ('사용자계정', 'USR_AUTH', '권한그룹', '경영그룹, 사용자그룹, 농장그룹, 물류현장그룹, 거래처그룹'),
    ('하우스', 'GHS_STATUS', '하우스상태', '휴경중, 생육중, 수확중, 정비중, 보수중'),
    ('하우스', 'GHS_USAGE', '하우스용도', '생산동, 육묘동, 관리동, 창고동'),
    ('설비', 'FAC_CAT', '설비분류', '관수, 난방, 환기, 조명, 기타'),
    ('농사장비', 'AGM_CAT', '장비분류', '경운기, 트랙터, 분무기, 예초기, 기타'),
    ('농사기구', 'AGT_CAT', '기구분류', '수확용, 정리용, 관리용, 기타'),
    ('비용관리', 'COST_CAT', '비용대분류', '인건비, 재료비, 경비, 감가상각, 기타'),
    ('비용관리', 'COST_DEPT', '적용부서', '전체, 농장, 온라인, 오프라인, 물류, 관리'),
    ('비용관리', 'COST_TAX', '세금구분', '계산서, 비계산서'),
    ('물류관리', 'LOG_VTYPE', '차량유형', '냉장탑, 화물, 승합, 기타'),
    ('배송타입', 'DLV_METHOD', '배송방법', '자체배송, 택배, 화물, 퀵'),
    ('판매부서', 'SLD_CH', '채널유형', '온라인, 오프라인'),
    ('판매처', 'SLR_PLAT', '플랫폼유형', '오픈마켓, 폐쇄몰, 직거래, 기타'),
    ('3정5S', 'TDO_CAT', '관리분류', '경영관리, 현장관리, 판매관리'),
    ('3정5S', 'TDO_CYCLE', '관리주기', '매일, 주간, 월간, 분기, 수시'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['사용 모듈', '그룹코드ID', '드롭다운 항목명', '현재 값']):
    hdr[i].text = h
for dd in dropdowns:
    row = table.add_row().cells
    for i, v in enumerate(dd):
        row[i].text = v

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run(f'총 {len(dropdowns)}개 그룹코드, 공통코드 관리 메뉴에서 통합 관리')
run.bold = True
run.font.color.rgb = RGBColor(26, 86, 219)

# ===== 4. 기능 설계 =====
doc.add_heading('4. 공통코드 관리 메뉴 기능 설계', level=1)

doc.add_heading('4-1. 메뉴 위치', level=2)
doc.add_paragraph('사이드바 최상단에 "공통코드 관리" 메뉴 추가', style='List Bullet')
doc.add_paragraph('전체 코드 관리보다 먼저 접근할 수 있도록 배치', style='List Bullet')

doc.add_heading('4-2. 화면 구성', level=2)

features = [
    ('그룹코드 목록', '좌측에 그룹코드 목록 표시\n클릭하면 해당 그룹의 상세코드 표시'),
    ('상세코드 관리', '우측에 선택된 그룹코드의 상세코드 목록\n신규 추가, 수정, 사용/정지 변경\n드래그로 정렬순서 변경'),
    ('검색 기능', '그룹코드명, 상세코드명으로 검색'),
    ('사용처 표시', '해당 코드가 어느 모듈에서 사용되는지 표시'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '기능'
hdr[1].text = '설명'
for feat, desc in features:
    row = table.add_row().cells
    row[0].text = feat
    row[1].text = desc

doc.add_heading('4-3. 입력 중 신규 코드 추가 기능', level=2)
doc.add_paragraph(
    '데이터 입력 중 드롭다운에 필요한 항목이 없을 때, '
    '드롭다운 옆의 [+] 버튼을 클릭하면 즉시 새 코드를 추가할 수 있습니다.'
)

steps = [
    '사용자가 마스터상품 등록 중 "대분류" 드롭다운을 봄',
    '필요한 분류가 없음 → 드롭다운 옆 [+신규] 버튼 클릭',
    '간편 입력창 팝업 (코드명만 입력)',
    '저장 → 드롭다운에 즉시 반영 → 선택 상태로 전환',
    '공통코드 관리에도 자동 등록',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('')

# ===== 5. 데이터 구조 =====
doc.add_heading('5. 데이터 구조 (DB 테이블)', level=1)

doc.add_heading('5-1. 그룹코드 테이블 (common_code_group)', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['필드명', '설명', '데이터타입', '비고']):
    hdr[i].text = h
group_fields = [
    ('group_code_id', '그룹코드ID', 'VARCHAR(30)', 'PK, 예: PRD_CAT_L'),
    ('group_code_name', '그룹코드명', 'VARCHAR(100)', '예: 마스터상품_대분류'),
    ('used_in_module', '사용 모듈', 'VARCHAR(50)', '예: master_product'),
    ('used_in_field', '사용 필드', 'VARCHAR(50)', '예: category_large'),
    ('description', '설명', 'VARCHAR(200)', ''),
    ('sort_order', '정렬순서', 'INT', ''),
    ('status', '사용상태', 'VARCHAR(10)', '사용/정지'),
    ('created_at', '등록일시', 'DATETIME', '자동'),
    ('updated_at', '수정일시', 'DATETIME', '자동'),
]
for f in group_fields:
    row = table.add_row().cells
    for i, v in enumerate(f):
        row[i].text = v

doc.add_heading('5-2. 상세코드 테이블 (common_code_detail)', level=2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['필드명', '설명', '데이터타입', '비고']):
    hdr[i].text = h
detail_fields = [
    ('detail_code_id', '상세코드ID', 'VARCHAR(40)', 'PK, 자동생성'),
    ('group_code_id', '그룹코드ID', 'VARCHAR(30)', 'FK → 그룹코드'),
    ('code_value', '코드값', 'VARCHAR(100)', '실제 드롭다운에 표시되는 값'),
    ('code_label_kh', '코드값(캄보디아어)', 'VARCHAR(100)', '다국어 지원'),
    ('sort_order', '정렬순서', 'INT', '드롭다운 표시 순서'),
    ('is_default', '기본값 여부', 'BOOLEAN', '드롭다운 기본선택'),
    ('status', '사용상태', 'VARCHAR(10)', '사용/정지'),
    ('created_at', '등록일시', 'DATETIME', '자동'),
    ('updated_at', '수정일시', 'DATETIME', '자동'),
]
for f in detail_fields:
    row = table.add_row().cells
    for i, v in enumerate(f):
        row[i].text = v

# ===== 6. 개발 진행 방안 =====
doc.add_heading('6. 개발 진행 방안 (추천)', level=1)

doc.add_heading('방안 A: 공통코드 메뉴 먼저 개발 (추천)', level=2)

plan_a = [
    ('1단계', '공통코드 관리 메뉴 개발', '공통코드 그룹/상세 CRUD\n좌측 그룹 목록 + 우측 상세코드 관리 화면'),
    ('2단계', '기존 드롭다운 데이터 이전', '현재 하드코딩된 34개 드롭다운 항목을\n공통코드 데이터로 자동 등록'),
    ('3단계', '기존 코드 24종 드롭다운 연동', '각 코드의 드롭다운을 공통코드 데이터에서\n동적으로 불러오도록 변경'),
    ('4단계', '입력 중 [+신규] 버튼 추가', '드롭다운 옆에 즉시 코드 추가 버튼 삽입'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['단계', '작업', '내용']):
    hdr[i].text = h
for step in plan_a:
    row = table.add_row().cells
    for i, v in enumerate(step):
        row[i].text = v

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('장점: ')
run.bold = True
doc.add_paragraph('기초정보 코드 입력 전에 필요한 드롭다운 값을 미리 준비 가능', style='List Bullet')
doc.add_paragraph('한 곳에서 전체 코드를 관리하므로 일관성 유지', style='List Bullet')
doc.add_paragraph('입력 중 코드 추가가 자연스럽게 동작', style='List Bullet')

doc.add_heading('방안 B: 기존 코드 입력하면서 동시 개발', level=2)
doc.add_paragraph('기존 24종 코드를 사용하면서 부족한 드롭다운이 나올 때마다 추가하는 방식')
p = doc.add_paragraph()
run = p.add_run('단점: ')
run.bold = True
doc.add_paragraph('코드 추가 요청이 빈번하여 작업 흐름이 끊김', style='List Bullet')
doc.add_paragraph('나중에 구조 변경 시 재작업 발생', style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('▶ 결론: 방안 A를 추천합니다.')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(26, 86, 219)

doc.add_paragraph(
    '공통코드 관리 메뉴를 먼저 개발하고, 필요한 드롭다운 값을 등록한 후, '
    '기초정보 24종 코드 입력을 시작하는 것이 가장 효율적입니다. '
    '이후 입력 과정에서 새로운 코드가 필요하면 [+신규] 버튼으로 즉시 추가할 수 있습니다.'
)

# ===== 7. 화면 와이어프레임 =====
doc.add_heading('7. 화면 구성 (와이어프레임)', level=1)

wireframe = """
┌─────────────────────────────────────────────────────────────────┐
│  GTP  그트플 통합관리시스템 - 공통코드 관리                        │
├─────────┬───────────────────────────────────────────────────────┤
│         │  [검색: ________________]  [+ 그룹코드 추가]            │
│ 공통코드 ├───────────────────────────────────────────────────────┤
│ 관리    │                                                       │
│         │  ┌──────────────────┐  ┌──────────────────────────┐   │
│ ─────── │  │ 그룹코드 목록     │  │ 상세코드 관리              │   │
│ 상품관리 │  │                  │  │                          │   │
│  마스터  │  │ ▸ 마스터상품_대분류│  │ 그룹: 마스터상품_대분류     │   │
│  SKU    │  │   마스터상품_중분류│  │ 사용: master_product      │   │
│  매핑   │  │   마스터상품_단위  │  │                          │   │
│ ─────── │  │   마스터상품_원산지│  │ [+ 상세코드 추가]          │   │
│ 자재관리 │  │   매입채널       │  │                          │   │
│  포장자재│  │   보관구분       │  │ 코드값    정렬  상태  작업  │   │
│  소모품  │  │   ...           │  │ ──────────────────────── │   │
│  물류자재│  │                  │  │ 엽채류     1   사용  ✏🗑  │   │
│ ─────── │  │ ▸ SKU상품_포장형태│  │ 과채류     2   사용  ✏🗑  │   │
│ 거래처  │  │   SKU상품_출고방식│  │ 근채류     3   사용  ✏🗑  │   │
│  ...    │  │   ...           │  │ 과일류     4   사용  ✏🗑  │   │
│         │  │                  │  │ 양념채소   5   사용  ✏🗑  │   │
│         │  └──────────────────┘  │ 버섯류     6   사용  ✏🗑  │   │
│         │                        │ 기타       7   사용  ✏🗑  │   │
│         │                        └──────────────────────────┘   │
└─────────┴───────────────────────────────────────────────────────┘

[입력 중 코드 추가 예시]
┌──────────────────────────┐
│ 대분류: [엽채류    ▼][+]  │  ← [+] 클릭시 아래 팝업
│                          │
│  ┌────────────────────┐  │
│  │ 신규 대분류 추가     │  │
│  │ 코드값: [________]  │  │
│  │ [취소] [저장]       │  │
│  └────────────────────┘  │
└──────────────────────────┘
"""

p = doc.add_paragraph()
run = p.add_run(wireframe)
run.font.name = 'Consolas'
run.font.size = Pt(8)

# 저장
output_path = r'C:\Users\User01\Desktop\GTP SYSTEM\GTP-SYSTEM\공통코드_관리체계_설계안.docx'
doc.save(output_path)
print(f'문서 생성 완료: {output_path}')
